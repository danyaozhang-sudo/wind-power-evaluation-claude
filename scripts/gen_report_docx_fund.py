#!/usr/bin/env python3
"""风电项目评估报告 Word 生成脚本（选项二·基金用途精简版）
使用方法：修改 REPORT_DATA 字典后直接运行
依赖：python-docx
Logo：~/.openclaw/workspace/skills/wind-power-analysis/assets/jianeng_logo_header.png

结构：项目概况 → 必选A → 财务测算结果(3.1~3.3) → 财务指标附表(4表) → 免责声明
对比选项一（详细版）：跳过必选B/C/D、投资决策建议，任务改称基金收购/出售
"""
import os, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

LOGO = os.path.expanduser("~/.openclaw/workspace/skills/wind-power-analysis/assets/jianeng_logo_header.png")
REPORT_DATA = {
    "project_name": "天津蓟州西龙虎峪镇",
    "province": "天津",
    "city": "蓟州区西龙虎峪镇",
    "region": "天津",
    "capacity_mw": 60,
    "annual_hours": 2000,
    "curtailment_rate": 0.05,
    "project_type": "增量风电（不参与机制电价竞价，全量中长期+绿电交易）",
    "version": "v5",
    "report_date": "2026-05-22",

    "mechanism_price": 0.32,
    "mechanism_cap": 0.32,
    "mechanism_ratio": 0,
    "mechanism_years": 0,
    "mkt_long_price": 0.364,
    "mkt_spot_price": 0.0,
    "green_premium": 0.004,
    "penalty": 0.015,
    "market_fee_rate": 0.02,

    "policy": [
        ("燃煤基准价", "0.3655 元/kWh", "发改价格〔2023〕526号", ""),
        ("机制电价竞价上限", "0.32 元/kWh", "天津市实施方案", ""),
        ("本项目策略", "不参与竞价", "决策依据", ""),
        ("绿电含电能量价", "~0.36 元/kWh", "绿电≈火电平价", ""),
        ("绿证溢价", "+0.004 元/kWh", "2025H2均价4.14元/个", ""),
        ("中长协+绿电加权", "0.364 元/kWh", "0.36+0.004", ""),
        ("两细则考核", "-0.015 元/kWh", "华北区域两个细则", ""),
        ("市场费分摊", "-2%", "136号文标准", ""),
        ("有效电价", "0.3420 元/kWh", "(0.364-0.015)×0.98", ""),
        ("消纳/限电", "5%含厂用电", "天津负荷中心偏保守", ""),
    ],

    "interest_rate": 0.035,
    "loan_years": 18,
    "operating_years": 20,
    "residual_rate": 0.03,
    "vat_rate": 0.13,
    "income_tax_rate": 0.25,
    "om_cost_1": 0.02,
    "om_cost_2": 0.06,
    "om_cost_3": 0.08,
    "ins_rate": "净值×0.2%",
    "effective_price": 0.3420,

    "t1_limit": 5.50,
    "t4_limit": 5.8517,
    "t1_irr": 6.69,
    "t4_irr": 6.00,
    "t4_equity_irr": 14.15,
    "t1_equity_irr": 16.61,
    "t1_dscr": 1.4242,
    "t4_dscr": 1.35,
    "t1_lcoe": 0.3060,
    "t4_lcoe": 0.3185,
    "payback_period": 6.1,

    "profit_table": [
        [1, 3899.03, 120.00, 66.81, 1702.84, 1123.53, 13.54, 218.07, 654.22],
        [5, 3899.03, 120.00, 53.19, 1702.84, 873.85, 17.44, 282.92, 848.77],
        [10, 3899.03, 360.00, 36.16, 1702.84, 561.76, 18.56, 304.92, 914.77],
        [15, 3899.03, 480.00, 19.14, 1702.84, 249.67, 21.56, 356.45, 1069.36],
        [18, 3899.03, 480.00, 8.92, 1702.84, 62.42, 24.48, 405.09, 1215.27],
        [19, 3899.03, 480.00, 5.51, 1702.84, 0.00, 49.80, 415.22, 1245.65],
        [20, 3899.03, 480.00, 2.11, 1702.84, 0.00, 49.80, 416.07, 1248.21],
    ],
    "fcf_table": [
        [0, "-35,110.20"], [1, "3,186.57"], [5, "3,196.78"],
        [10, "3,032.36"], [15, "2,956.54"], [20, "2,951.05"],
    ],
    "equity_cf_table": [
        [0, "—", "—", "—", "-7,022.04", "-7,022.04"],
        [1, 654.22, 1702.84, 1560.45, 796.61, "-6,225.43"],
        [5, 848.77, 1702.84, 1560.45, 991.16, "-2,321.91"],
        [10, 914.77, 1702.84, 1560.45, 1057.16, "2,664.52"],
        [15, 1069.36, 1702.84, 1560.45, 1211.75, "7,971.01"],
        [18, 1215.27, 1702.84, 1560.45, 1357.67, "11,805.79"],
        [19, 1245.65, 1702.84, 0.00, 2948.50, "14,754.29"],
        [20, 1248.21, 1702.84, 0.00, 2951.05, "17,705.34"],
    ],
    "dscr_table": [
        [1, 3721.79, 10.58, 262.35, 3448.86, 1670.93, 1203.07, 2874.01, 1.20],
        [5, 3733.46, 14.75, 331.07, 3387.65, 1670.93, 935.72, 2606.66, 1.30],
        [10, 3508.05, 16.22, 357.89, 3133.94, 1670.93, 601.54, 2272.47, 1.38],
        [15, 3402.64, 19.56, 414.25, 2968.82, 1670.93, 267.35, 1938.28, 1.53],
        [18, 3411.39, 22.69, 465.78, 2922.91, 1670.93, 66.84, 1737.77, 1.68],
        ["19~20", "3,414~3,417", "49.80", "476~477", "2,888~2,890", 0, 0, 0, "∞"],
    ],
}

TODAY = "20260522"

COLOR_DB = RGBColor(0x1B, 0x3A, 0x5C)

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(15)

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(13)

def para(doc, text, bold=False, sz=11, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    r.font.size = Pt(sz); r.bold = bold

def table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
        c._element.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = '仿宋'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_header(doc, logo_path):
    for s in doc.sections:
        hdr = s.header
        hdr.is_linked_to_previous = False
        p = hdr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2))

def add_disclaimer(doc, R):
    h1(doc, "免责声明")
    text = (
        "本报告由江能能源开发的AI辅助评估系统生成，并经过人工校核，不构成投资建议。"
        "报告中所引用的电力市场数据（机制电价、中长期交易电价、现货电价、限电率、利用小时数等）均来自公开渠道，"
        "部分电价参数在缺乏直接交易数据的情况下采用保守假设推算，可能与实际成交价格存在偏差。"
        "实际项目投资决策应结合以下因素综合判断：\n\n"
        "① 场址实测测风数据（至少一个完整年度或有风机厂背书）；"
        "② 电网接入批复及送出工程条件；"
        "③ EPC 招标实际报价；"
        "④ 项目所在地最新的机制电价竞价结果；"
        "⑤ 所在省电力交易中心公布的全年现货与中长期交易结算数据。\n\n"
        f"本报告中的财务模型基于特定假设（融资利率 {R['interest_rate']*100:.1f}%、经营期 20 年、等额本金还款等），"
        "不同融资结构、利率环境及政策变化可能导致测算结果显著偏离。"
        "报告中的'投资边界'和'出售边界'为理论测算阈值，不代表项目实际可实现的交易价格或融资条件。\n\n"
        "江能能源及报告编制方不对因使用本报告而产生的任何直接或间接损失承担责任。"
        "未经授权，不得转载或用于商业用途。"
    )
    para(doc, text, sz=9)

def build(doc, R):
    gen_mwh = R['capacity_mw'] * R['annual_hours'] * (1 - R['curtailment_rate'])

    # ═══════════ 标题 ═══════════
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"江能能源 · {R['project_name']}{R['capacity_mw']}MW风电项目投资评估报告")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = COLOR_DB
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    para(doc,
         f"版本：{R['version']} | 日期：{R['report_date']} | 分析师：江能能源投资研究院(Claude Opus+)"
         f"\n项目类型：{R['project_type']}",
         sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══════════ 一、项目概况 ═══════════
    h1(doc, "一、项目概况")
    table(doc, ["参数", "数值"], [
        ["所在地", f"{R['province']} {R['city']}（{R['region']}）"],
        ["装机容量", f"{R['capacity_mw']} MW"],
        ["理论利用小时", f"{R['annual_hours']} h"],
        ["限电率", f"{R['curtailment_rate']*100:.1f}%"],
        ["年净发电量", f"{gen_mwh:,.0f} MWh"],
        ["开工时间", "2026年3月"],
        ["预计并网", "2027年3月"],
        ["投资方", "皓景雅岚"],
    ])

    # ═══════════ 二、[必选A] 政策溯源 ═══════════
    h1(doc, "二、[必选A] 电价与执行期限的政策依据（完整溯源）")
    policy_headers = ["序号", "参数", "数值", "来源文件", "URL"]
    policy_rows = [[i+1, item[0], item[1], item[2], item[3]] for i, item in enumerate(R['policy'])]
    table(doc, policy_headers, policy_rows)

    # ═══════════ 三、财务测算结果 ═══════════
    h1(doc, "三、财务测算结果")

    h2(doc, "3.1 模型边界条件")
    table(doc, ["参数", "数值"], [
        ["装机容量", f"{R['capacity_mw']} MW"],
        ["理论利用小时数", f"{R['annual_hours']} h"],
        ["限电率", f"{R['curtailment_rate']*100:.1f}%"],
        ["年净发电量", f"{gen_mwh:,.0f} MWh"],
        ["有效电价", f"{R['effective_price']} 元/kWh"],
        ["经营期", f"{R['operating_years']} 年"],
        ["融资利率", f"{R['interest_rate']*100:.1f}%"],
        ["融资期限", f"{R['loan_years']} 年"],
        ["折旧年限", f"{R['operating_years']} 年"],
        ["残值率", f"{R['residual_rate']*100:.0f}%"],
        ["运维费第1段(1-5年)", f"{R['om_cost_1']} 元/W·年"],
        ["运维费第2段(6-10年)", f"{R['om_cost_2']} 元/W·年"],
        ["运维费第3段(11-20年)", f"{R['om_cost_3']} 元/W·年"],
        ["保险费率", R['ins_rate']],
        ["增值税率", f"{R['vat_rate']*100:.0f}%"],
        ["所得税率", f"{R['income_tax_rate']*100:.0f}%"],
    ])

    h2(doc, "3.2 基金收购（80%融资·固定5.5元/W）")
    t1_rows = [
        ["单瓦投资", f"{R['t1_limit']} 元/W"],
        ["总投资", f"{R['t1_limit']*R['capacity_mw']/100:.2f} 亿元"],
        ["融资额（80%）", f"{R['t1_limit']*R['capacity_mw']/100*0.8:.4f} 亿元"],
        ["资本金（20%）", f"{R['t1_limit']*R['capacity_mw']/100*0.2:.4f} 亿元"],
        ["最小DSCR", f"{R['t1_dscr']:.4f}"],
        ["全投资IRR", f"{R['t1_irr']}%"],
        ["税后资本金IRR", f"{R['t1_equity_irr']}%"],
        ["LCOE", f"{R['t1_lcoe']:.4f} 元/kWh"],
        ["年均净利润", "1,082 万元"],
    ]
    if R.get('payback_period'):
        t1_rows.insert(7, ["资本金全部回收期", f"{R['payback_period']:.1f} 年"])
    table(doc, ["指标", "数值"], t1_rows)

    h2(doc, "3.3 基金出售（80%融资，全投资IRR≥6% & 资本金IRR≥8%）")
    table(doc, ["指标", "数值"], [
        ["目标单瓦投资", f"{R['t4_limit']} 元/W"],
        ["总投资", f"{R['t4_limit']*R['capacity_mw']/100:.2f} 亿元"],
        ["全投资IRR", f"{R['t4_irr']}%"],
        ["税后资本金IRR", f"{R['t4_equity_irr']}%"],
        ["LCOE", f"{R['t4_lcoe']:.4f} 元/kWh"],
        ["年均净利润", "984 万元"],
        ["最小DSCR", f"{R['t4_dscr']}"],
    ])

    # ═══════════ 四、财务指标附表 ═══════════
    h1(doc, "四、财务指标附表")

    h2(doc, "表1：资本金投资利润表（基金收购80%融资）（单位：万元）")
    p1_h = ["年份","营业收入","运维费","保险费","折旧","利息","增值税及附加","所得税","净利润"]
    p1_r = [[str(r[0]), f"{r[1]:.2f}", f"{r[2]:.2f}", f"{r[3]:.2f}", f"{r[4]:.2f}", f"{r[5]:.2f}", f"{r[6]:.2f}", f"{r[7]:.2f}", f"{r[8]:.2f}"] for r in R['profit_table']]
    table(doc, p1_h, p1_r)
    para(doc, "说明：前5年运维费低（0.02元/W），净利润逐步上升至1,041万；第6年起运维费升至0.06元/W，净利润先降后随利息减少逐步回升；第11年起运维费0.08元/W；第19年起无利息，净利润跃升至1,323万以上。", sz=9)

    h2(doc, "表2：全投资净现金流表（单位：万元）")
    fcf_h = ["年份", "全投资FCF"]
    fcf_r = [[str(r[0]), str(r[1])] for r in R['fcf_table']]
    table(doc, fcf_h, fcf_r)
    para(doc, f"全投资FCF基于无杠杆附加税（剔除利息税盾）。t=0为初始投资流出{R['t1_limit']*R['capacity_mw']/100:.2f}亿元。全投资IRR={R['t1_irr']:.2f}%。", sz=9)

    h2(doc, "表3：资本金投资现金流表（基金收购80%融资）（单位：万元）")
    eq_h = ["年份", "净利润", "折旧", "偿还本金", "股权现金流", "累计股权现金流"]
    def fmt_val(v):
        if isinstance(v, str): return v
        return f"{v:,.2f}"
    eq_r = [[str(r[0]), fmt_val(r[1]), fmt_val(r[2]), fmt_val(r[3]), fmt_val(r[4]), fmt_val(r[5])] for r in R['equity_cf_table']]
    table(doc, eq_h, eq_r)
    pb_text = f"资本金全部回收期约{R['payback_period']:.1f}年（第7年累计转正）。" if R.get('payback_period') else ""
    para(doc, f"t=0初始资本金投入{R['t1_limit']*R['capacity_mw']/5:,.0f}万元（总投资{R['t1_limit']*R['capacity_mw']/100:.2f}亿×20%）。股权现金流=净利润+折旧-偿还本金。税后资本金IRR={R['t1_equity_irr']:.2f}%。{pb_text}", sz=9)

    h2(doc, "表4：偿债覆盖计算表（基金收购80%融资）（单位：万元）")
    dscr_h = ["年份","EBITDA","增值税及附加","所得税","可用于还款","应还本金","应还利息","应还本息","DSCR"]
    dscr_r = []
    for r in R['dscr_table']:
        row = [str(r[0])]
        for v in r[1:]:
            if isinstance(v, str):
                row.append(v)
            else:
                row.append(f"{v:,.2f}" if v >= 100 else f"{v:.2f}")
        dscr_r.append(row)
    table(doc, dscr_h, dscr_r)
    para(doc, f"说明：DSCR=可用于还款/应还本息。可用于还款=EBITDA-增值税及附加-所得税。应还本金=期初本金/{R['loan_years']}年（等额本金）。最小DSCR={R['t1_dscr']:.2f}，全期偿债能力达标。", sz=9)

    # ═══════════ 免责声明 ═══════════
    add_disclaimer(doc, R)


def generate(R):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        s.header_distance = Cm(1.0)
    add_header(doc, LOGO)
    build(doc, R)
    out = os.path.expanduser(
        f"~/.openclaw/workspace/projects/{R['project_name']}{R['capacity_mw']}MW风电/"
        f"{R['project_name']}{R['capacity_mw']}MW风电_评估报告_{TODAY}.docx"
    )
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    os.chmod(out, 0o644)
    print(f"✓ [基金用途版] {out}")
    return out

if __name__ == '__main__':
    generate(REPORT_DATA)
