#!/usr/bin/env python3
"""风电项目评估报告 Word 生成脚本（详细版）
使用方法：修改 REPORT_DATA 字典后直接运行
依赖：python-docx (pip install python-docx --break-system-packages)
Logo：~/.hermes/skills/openclaw-imports/wind-power-analysis/assets/jianeng_logo_header.png

版本匹配：report_template.md（风电详细结构）
"""
import os, sys, io, tempfile
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image
from lxml import etree

# ───────────────────────────────────────────────
# >>> 配置（每个项目修改此处）<<<
# ───────────────────────────────────────────────
LOGO = os.path.expanduser("~/.openclaw/workspace/skills/wind-power-analysis/assets/jianeng_logo_header.png")
REPORT_DATA = {
    # ── 项目基本信息 ──
    "project_name": "天津蓟州西龙虎峪镇",
    "province": "天津",
    "city": "蓟州区西龙虎峪镇",
    "region": "天津",
    "capacity_mw": 60,
    "annual_hours": 2000,
    "curtailment_rate": 0.05,
    "curtailment_source": "厂商实测数据（2026年）",
    "curtailment_detail": "天津位于负荷中心，5%偏保守（含厂用电）",
    "project_type": "增量风电（不参与机制电价竞价，全量中长期+绿电）",
    "version": "v4",
    "report_date": "2026-05-22",

    # ── 电价参数 ──
    "mechanism_price": 0.32,        # 竞价上限（含税）
    "mechanism_cap": 0.32,          # 竞价上限
    "mechanism_ratio": 0,           # 不参与
    "mechanism_years": 0,           # 不适用
    "base_hours": 0,                # 不适用
    "mkt_long_price": 0.364,        # 中长协+绿电加权价
    "mkt_spot_price": 0.0,          # 天津无现货
    "green_premium": 0.004,         # 绿证溢价
    "penalty": 0.015,               # 两细则考核
    "market_fee_rate": 0.02,        # 市场费用分摊

    # ── 政策溯源10项（[必选A]）──
    "policy": [
        ("燃煤基准价", "0.3655 元/kWh", "发改价格〔2023〕526号", ""),
        ("机制电价竞价上限", "0.32 元/kWh", "天津市实施方案", ""),
        ("本项目策略", "不参与竞价", "决策依据", ""),
        ("绿电含电能量价", "~0.36 元/kWh", "绿电≈火电平价", ""),
        ("绿证溢价", "+0.004 元/kWh", "2025H2均价4.14元/个", ""),
        ("中长协+绿电加权", "0.364 元/kWh", "0.36+0.004", ""),
        ("两细则考核", "-0.015 元/kWh", "华北区域两个细则", ""),
        ("市场费分摊", "-2%", "136号文标准", ""),
        ("有效电价", "0.3420 元/kWh", "(0.364-0.015)×0.98", ""),
        ("消纳/限电", "5%含厂用电", "天津负荷中心偏保守", ""),
    ],

    # ── 计算过程参数 ──
    "calc_method": "单一电价法（天津无现货不参与机制竞价）",
    "calc_method_note": "天津无现货市场，不参与机制电价竞价，全电量中长期+绿电，单一电价20年不变",
    "calc_mechanism_before": "N/A",
    "calc_mechanism_after": "N/A",
    "calc_market_before": "0.364 元/kWh（绿电0.36+绿证0.004）",
    "calc_market_after": "0.3420 元/kWh = (0.364-0.015)×0.98",

    # ── 13行多口径表（[必选C]）──
    "pricing_rows": [
        ("①", "燃煤基准价（含税）", 0.3655, "锚点", "❌"),
        ("②", "机制电价竞价上限", 0.320, "不参与", "❌"),
        ("③", "绿电含电能量价", 0.360, "风电长期绿电价", "✅"),
        ("④", "绿证溢价", 0.004, "4元/个", "✅"),
        ("⑤", "中长协+绿电加权（扣费前）", 0.364, "中间值", "❌"),
        ("⑥", "两细则考核费", -0.015, "华北规则", "✅"),
        ("⑦", "市场费分摊", -0.0073, "2%", "✅"),
        ("⑧", "两细则后电价", 0.349, "中间值", "❌"),
        ("⑨", "有效电价（扣费后含税）", 0.3420, "模型输入", "✅"),
        ("⑩", "有效电价（不含税）", 0.3027, "⑨÷1.13", "✅"),
        ("⑪", "机制电价中标价参考", 0.320, "未参与", "❌"),
        ("⑫", "差额电量电价", 0.000, "N/A", "N/A"),
        ("⑬", "全生命周期加权均价", 0.3420, "20年不变", "❌"),
    ],

    # ── 敏感性分析（[必选D]）──
    "sensitivity": {
        "conservative": {"basis": "有效电价0.33元", "spot": 0.33, "eff_price": 0.330, "t1": 4.78, "t4": 5.60},
        "neutral": {"basis": "有效电价0.342元", "spot": 0.34, "eff_price": 0.3420, "t1": 5.01, "t4": 5.85},
        "optimistic": {"basis": "有效电价0.36元", "spot": 0.36, "eff_price": 0.360, "t1": 5.28, "t4": 6.16},
    },

    # ── 模型边界条件 ──
    "interest_rate": 0.04,
    "loan_years": 18,
    "operating_years": 20,
    "residual_rate": 0.03,
    "vat_rate": 0.13,
    "income_tax_rate": 0.25,
    "om_cost_1": 0.02,
    "om_cost_2": 0.06,
    "om_cost_3": 0.08,
    "ins_rate": "净值×0.2%",
    "effective_price": 0.3420,

    # ── 测算结果 ──
    "t1_limit": 5.01,
    "t4_limit": 5.85,
    "t1_irr": 7.78,
    "t4_irr": 6.00,
    "t4_equity_irr": 13.14,
    "t1_dscr": 1.2000,
    "t4_dscr": 1.30,
    "t1_lcoe": 0.2836,
    "t4_lcoe": 0.3167,

    # ── 投资决策建议 ──
    "rec_sell_limit": 5.85,
    "rec_buy_limit": 5.01,
    "risk_notes": "①投资协议2026年6月到期②环评批复进度③中长协定价波动④利率4%",
    "rating": "★★★★☆",

    # ── 第八节·财务附表数据（关键年份）──
    "profit_table": [
        [1, 3899.03, 120.00, 66.81, 1702.84, 1123.53, 13.54, 218.07, 654.22],
        [5, 3899.03, 120.00, 53.19, 1702.84, 873.85, 17.44, 282.92, 848.77],
        [10, 3899.03, 360.00, 36.16, 1702.84, 561.76, 18.56, 304.92, 914.77],
        [15, 3899.03, 480.00, 19.14, 1702.84, 249.67, 21.56, 356.45, 1069.36],
        [18, 3899.03, 480.00, 8.92, 1702.84, 62.42, 24.48, 405.09, 1215.27],
        [19, 3899.03, 480.00, 5.51, 1702.84, 0.00, 49.80, 415.22, 1245.65],
        [20, 3899.03, 480.00, 2.11, 1702.84, 0.00, 49.80, 416.07, 1248.21],
    ],
    "fcf_table": [
        [0, "-35,110.20"], [1, "3,186.57"], [5, "3,196.78"],
        [10, "3,032.36"], [15, "2,956.54"], [20, "2,951.05"],
    ],
    "equity_cf_table": [
        [0, "—", "—", "—", "-7,022.04"],
        [1, 654.22, 1702.84, 1560.45, 796.61],
        [5, 848.77, 1702.84, 1560.45, 991.16],
        [10, 914.77, 1702.84, 1560.45, 1057.16],
        [15, 1069.36, 1702.84, 1560.45, 1211.75],
        [18, 1215.27, 1702.84, 1560.45, 1357.67],
        [19, 1245.65, 1702.84, 0.00, 2948.50],
        [20, 1248.21, 1702.84, 0.00, 2951.05],
    ],
    "dscr_table": [
        [1, 3721.79, 10.58, 262.35, 3448.86, 1670.93, 1203.07, 2874.01, 1.20],
        [5, 3733.46, 14.75, 331.07, 3387.65, 1670.93, 935.72, 2606.66, 1.30],
        [10, 3508.05, 16.22, 357.89, 3133.94, 1670.93, 601.54, 2272.47, 1.38],
        [15, 3402.64, 19.56, 414.25, 2968.82, 1670.93, 267.35, 1938.28, 1.53],
        [18, 3411.39, 22.69, 465.78, 2922.91, 1670.93, 66.84, 1737.77, 1.68],
        ["19~20", "3,414~3,417", "49.80", "476~477", "2,888~2,890", 0, 0, 0, "∞"],
    ],
}

TODAY = "20260522"
# ───────────────────────────────────────────────

COLOR_DB = RGBColor(0x1B, 0x3A, 0x5C)

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(15)

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(13)

def h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.color.rgb = COLOR_DB; r.font.name = '黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体'); r.font.size = Pt(11)

def para(doc, text, bold=False, sz=11, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.font.name = '仿宋'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    r.font.size = Pt(sz); r.bold = bold

def code_para(doc, text, sz=9):
    """等宽字体段落，用于展示计算过程"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(sz)

def table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
                r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C" w:val="clear"/>')
        c._element.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = '仿宋'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

def add_header(doc, logo_path):
    """江能 Logo 页眉，2cm 宽居右"""
    for s in doc.sections:
        hdr = s.header
        hdr.is_linked_to_previous = False
        p = hdr.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2))

def add_watermark(doc, logo_path, opacity=0.12):
    """在文档所有页面添加半透明Logo水印（防盗版）"""
    # 1. 生成半透明水印图片
    img = Image.open(logo_path)
    rgba = img.convert('RGBA')
    rgba = rgba.resize((360, 300), Image.LANCZOS)
    r, g, b, a = rgba.split()
    a = a.point(lambda x: min(x, int(255 * opacity)))
    watermark = Image.merge('RGBA', (r, g, b, a))
    tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    watermark.save(tmp.name, 'PNG')
    tmp_path = tmp.name

    # 2. 插入每个章节页眉，浮于文字下方
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(tmp_path, width=Cm(7))

        for drawing in run._element.findall('.//' + qn('w:drawing')):
            inline = drawing.find(qn('wp:inline'))
            if inline is not None:
                extent = inline.find(qn('wp:extent'))
                cx = extent.get('cx', '0') if extent is not None else '0'
                cy = extent.get('cy', '0') if extent is not None else '0'

                anchor = etree.SubElement(drawing, qn('wp:anchor'))
                anchor.set('simplePos', '0')
                anchor.set('relativeHeight', '0')
                anchor.set('behindDoc', '1')
                anchor.set('locked', '0')
                anchor.set('layoutInCell', '1')
                anchor.set('allowOverlap', '1')

                sp = etree.SubElement(anchor, qn('wp:simplePos'))
                sp.set('x', '0'); sp.set('y', '0')

                hp = etree.SubElement(anchor, qn('wp:positionH'))
                hp.set('relativeFrom', 'page')
                etree.SubElement(hp, qn('wp:align')).text = 'center'

                vp = etree.SubElement(anchor, qn('wp:positionV'))
                vp.set('relativeFrom', 'page')
                etree.SubElement(vp, qn('wp:align')).text = 'center'

                a_ext = etree.SubElement(anchor, qn('wp:extent'))
                a_ext.set('cx', cx); a_ext.set('cy', cy)

                ee = etree.SubElement(anchor, qn('wp:effectExtent'))
                ee.set('l', '0'); ee.set('t', '0'); ee.set('r', '0'); ee.set('b', '0')

                etree.SubElement(anchor, qn('wp:wrapNone'))

                for child in list(inline):
                    tag_local = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag_local in ('graphic', 'extent'):
                        anchor.append(child)

                drawing.remove(inline)

    try: os.unlink(tmp_path)
    except: pass


def add_disclaimer(doc):
    """免责声明 — 必须包含"""
    h1(doc, "免责声明")
    text = (
        "本报告由江能能源开发的AI辅助评估系统生成，并经过人工校核，不构成投资建议。"
        "报告中所引用的电力市场数据（机制电价、中长期交易电价、现货电价、限电率、利用小时数等）均来自公开渠道，"
        "部分电价参数在缺乏直接交易数据的情况下采用保守假设推算，可能与实际成交价格存在偏差。"
        "实际项目投资决策应结合以下因素综合判断：\n\n"
        "① 场址实测测风数据（至少一个完整年度）；"
        "② 电网接入批复及送出工程条件；"
        "③ EPC 招标实际报价；"
        "④ 项目所在地最新的机制电价竞价结果；"
        "⑤ 所在省电力交易中心公布的全年现货与中长期交易结算数据。\n\n"
        "本报告中的财务模型基于特定假设（融资利率 4%、经营期 20 年、等额本金还款等），"
        "不同融资结构、利率环境及政策变化可能导致测算结果显著偏离。"
        "报告中的'投资边界'和'出售边界'为理论测算阈值，不代表项目实际可实现的交易价格或融资条件。\n\n"
        "江能能源及报告编制方不对因使用本报告而产生的任何直接或间接损失承担责任。"
        "未经授权，不得转载或用于商业用途。"
    )
    para(doc, text, sz=9)

def build(doc, R):
    """构建详细报告正文（匹配 report_template.md 风电详细版）"""
    gen_mwh = R['capacity_mw'] * R['annual_hours'] * (1 - R['curtailment_rate'])

    # ═══════════ 标题 ═══════════
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(f"江能能源 · {R['project_name']}{R['capacity_mw']}MW风电项目投资评估报告")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = COLOR_DB
    r.font.name = '黑体'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    para(doc,
         f"版本：{R['version']} | 日期：{TODAY} | 分析师：江能研究院（Claude Opus辅助）"
         f"\n项目类型：{R['project_type']}",
         sz=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══════════ 一、项目概况 ═══════════
    h1(doc, "一、项目概况")
    table(doc, ["参数", "数值"], [
        ["所在地", f"{R['province']} {R['city']}（{R['region']}）"],
        ["装机容量", f"{R['capacity_mw']} MW"],
        ["理论利用小时", f"{R['annual_hours']} h"],
        ["限电率", f"{R['curtailment_rate']*100:.1f}%"],
        ["年净发电量", f"{gen_mwh:,.0f} MWh"],
    ])

    # ═══════════ 二、[必选A] 政策溯源 ═══════════
    h1(doc, "二、[必选A] 电价与执行期限的政策依据（完整溯源）")
    policy_headers = ["序号", "参数", "数值", "来源文件", "URL"]
    policy_rows = [[i+1, item[0], item[1], item[2], item[3]] for i, item in enumerate(R['policy'])]
    table(doc, policy_headers, policy_rows)

    # ═══════════ 三、[必选B] 计算过程 ═══════════
    h1(doc, "三、[必选B] 电价确认的计算过程")

    h2(doc, "第一步：原始参数确认")
    table(doc, ["参数", "数值", "来源"], [
        ["机制电价竞价上限", f"{R['mechanism_cap']}", R['policy'][0][2]],
        ["实际中标出清价", f"{R['mechanism_price']}", R['policy'][1][2]],
        ["机制电量比例", f"{R['mechanism_ratio']*100:.0f}%", R['policy'][2][2]],
        ["执行期限", f"{R['mechanism_years']}年", R['policy'][3][2]],
        ["现货均价（含税）", f"{R['mkt_spot_price']}", R['policy'][5][2]],
        ["中长期均价（含税）", f"{R['mkt_long_price']}", R['policy'][6][2]],
    ])

    h2(doc, "第二步：计算口径确认")
    para(doc, f"{R['province']}采用{R['calc_method']}：{R['calc_method_note']}", sz=10)

    h2(doc, "第三步：机制期电价")
    code_para(doc, f"扣费前 = {R['calc_mechanism_before']}")
    code_para(doc, f"有效电价 = {R['calc_mechanism_after']}")

    h2(doc, "第四步：市场化期电价")
    code_para(doc, f"扣费前 = {R['calc_market_before']}")
    code_para(doc, f"有效电价 = {R['calc_market_after']}")

    # ═══════════ 四、[必选C] 多口径表 ═══════════
    h1(doc, "四、[必选C] 电价多口径一览表（13行）")
    table(doc, ["序号", "电价口径", "数值(元/kWh)", "含义", "是否进模型"],
          R['pricing_rows'])

    # ═══════════ 五、[必选D] 敏感性 ═══════════
    h1(doc, "五、[必选D] 现货电价敏感性分析")
    table(doc, ["现货价假设", "取值依据", "机制期有效电价", "投资边界", "出售边界"], [
        ["保守", R['sensitivity']['conservative']['basis'],
         f"{R['sensitivity']['conservative']['eff_price']}", f"{R['sensitivity']['conservative']['t1']}", f"{R['sensitivity']['conservative']['t4']}"],
        ["★中性", R['sensitivity']['neutral']['basis'],
         f"{R['sensitivity']['neutral']['eff_price']}", f"{R['sensitivity']['neutral']['t1']}", f"{R['sensitivity']['neutral']['t4']}"],
        ["乐观", R['sensitivity']['optimistic']['basis'],
         f"{R['sensitivity']['optimistic']['eff_price']}", f"{R['sensitivity']['optimistic']['t1']}", f"{R['sensitivity']['optimistic']['t4']}"],
    ])

    # ═══════════ 六、测算结果 ═══════════
    # ═══════════ 六、财务测算结果 ═══════════
    h1(doc, "六、财务测算结果")

    h2(doc, "6.1 模型边界条件")
    table(doc, ["参数", "数值"], [
        ["装机容量", f"{R['capacity_mw']} MW"],
        ["理论利用小时数", f"{R['annual_hours']} h"],
        ["限电率", f"{R['curtailment_rate']*100:.1f}%"],
        ["年净发电量", f"{gen_mwh:,.0f} MWh"],
        ["有效电价", f"{R['effective_price']} 元/kWh"],
        ["经营期", f"{R['operating_years']} 年"],
        ["融资利率", f"{R['interest_rate']*100:.0f}%"],
        ["融资期限", f"{R['loan_years']} 年"],
        ["折旧年限", f"{R['operating_years']} 年"],
        ["残值率", f"{R['residual_rate']*100:.0f}%"],
        ["运维费第1段(1-5年)", f"{R['om_cost_1']} 元/W·年"],
        ["运维费第2段(6-10年)", f"{R['om_cost_2']} 元/W·年"],
        ["运维费第3段(11-20年)", f"{R['om_cost_3']} 元/W·年"],
        ["保险费率", R['ins_rate']],
        ["增值税率", f"{R['vat_rate']*100:.0f}%"],
        ["所得税率", f"{R['income_tax_rate']*100:.0f}%"],
    ])

    h2(doc, "6.2 任务1：投资边界（100%融资，DSCR≥1.2）")
    table(doc, ["指标", "数值"], [
        ["最高单瓦投资", f"{R['t1_limit']} 元/W"],
        ["总投资", f"{R['t1_limit']*R['capacity_mw']/100:.2f} 亿元"],
        ["最小DSCR", f"{R['t1_dscr']:.4f}"],
        ["全投资IRR", f"{R['t1_irr']}%"],
        ["LCOE", f"{R['t1_lcoe']:.4f} 元/kWh"],
    ])

    h2(doc, "6.3 任务4：出售边界（80%融资，IRR≥6%+资本金IRR≥8%）")
    table(doc, ["指标", "数值"], [
        ["目标单瓦投资", f"{R['t4_limit']} 元/W"],
        ["总投资", f"{R['t4_limit']*R['capacity_mw']/100:.2f} 亿元"],
        ["全投资IRR", f"{R['t4_irr']}%"],
        ["资本金IRR", f"{R['t4_equity_irr']}%"],
        ["LCOE", f"{R['t4_lcoe']:.4f} 元/kWh"],
    ])

    # ═══════════ 七、投资决策建议 ═══════════
    h1(doc, "七、投资决策建议")
    table(doc, ["建议项", "内容"], [
        ["推荐单瓦投资阈值", f"{R['rec_sell_limit']}元/W（出售边界）"],
        ["投资边界（保底门槛）", f"{R['rec_buy_limit']}元/W"],
        ["风险提示", R['risk_notes']],
        ["综合评价", R['rating']],
    ])

    # ═══════════ 八、财务指标附表 ═══════════
    h1(doc, "八、财务指标附表")

    h2(doc, "表1：资本金投资利润表（出售边界·80%融资）（单位：万元）")
    p1_h = ["年份","营业收入","运维费","保险费","折旧","利息","增值税及附加","所得税","净利润"]
    p1_r = [[str(r[0]), f"{r[1]:.2f}", f"{r[2]:.2f}", f"{r[3]:.2f}", f"{r[4]:.2f}", f"{r[5]:.2f}", f"{r[6]:.2f}", f"{r[7]:.2f}", f"{r[8]:.2f}"] for r in R['profit_table']]
    table(doc, p1_h, p1_r)
    para(doc, "说明：营业收入=年发电量×有效电价 | 运维费=装机×单位费率(分段) | 保险费=净值×0.2% | 折旧=(总投资×(1-残值率))/20年 | 利息=期初余额×利率 | 增值税=销项税-进项税 | 所得税=(收入-运维-保险-折旧-利息-增值税及附加)×25% | 净利润=营业收入-运维费-保险费-折旧-利息-增值税及附加-所得税", sz=9)

    h2(doc, "表2：全投资净现金流表（出售边界口径）（单位：万元）")
    fcf_h = ["年份", "全投资FCF"]
    fcf_r = [[str(r[0]), str(r[1])] for r in R['fcf_table']]
    table(doc, fcf_h, fcf_r)
    para(doc, f"全投资FCF基于无杠杆附加税（剔除利息税盾）。t=0为初始投资流出{R['t4_limit']*R['capacity_mw']/100:.2f}亿元。全投资IRR={R['t4_irr']:.2f}%。", sz=9)

    h2(doc, "表3：资本金投资现金流表（出售边界·80%融资）（单位：万元）")
    eq_h = ["年份", "净利润", "折旧", "偿还本金", "股权现金流"]
    def fmt_val(v):
        if isinstance(v, str): return v
        return f"{v:,.2f}"
    eq_r = [[str(r[0]), fmt_val(r[1]), fmt_val(r[2]), fmt_val(r[3]), fmt_val(r[4])] for r in R['equity_cf_table']]
    table(doc, eq_h, eq_r)
    para(doc, f"t=0初始资本金投入{R['t4_limit']*R['capacity_mw']/5:,.0f}万元（总投资{R['t4_limit']*R['capacity_mw']/100:.2f}亿×20%）。股权现金流=净利润+折旧-偿还本金。税后资本金IRR={R['t4_equity_irr']:.2f}%。", sz=9)

    h2(doc, "表4：偿债覆盖计算表（投资边界·100%融资）（单位：万元）")
    dscr_h = ["年份","EBITDA","增值税及附加","所得税","可用于还款","应还本金","应还利息","应还本息","DSCR"]
    dscr_r = []
    for r in R['dscr_table']:
        row = [str(r[0])]
        for v in r[1:]:
            if isinstance(v, str):
                row.append(v)
            else:
                row.append(f"{v:,.2f}" if v >= 100 else f"{v:.2f}")
        dscr_r.append(row)
    table(doc, dscr_h, dscr_r)
    para(doc, f"说明：DSCR=可用于还款/应还本息。可用于还款=EBITDA-增值税及附加-所得税。应还本金=期初本金/{R['loan_years']}年（等额本金）。最小DSCR={R['t1_dscr']:.2f}，全期偿债能力达标。", sz=9)

    # ═══════════ 免责声明 ═══════════
    add_disclaimer(doc)


def generate(R):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
        s.header_distance = Cm(1.0)
    add_header(doc, LOGO)
    build(doc, R)
    add_watermark(doc, LOGO)
    out = os.path.expanduser(
        f"~/.openclaw/workspace/projects/{R['project_name']}{R['capacity_mw']}MW风电/"
        f"{R['project_name']}{R['capacity_mw']}MW风电_评估报告_{TODAY}.docx"
    )
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f"✓ {out}")

if __name__ == '__main__':
    generate(REPORT_DATA)
